import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH

nameTxt = open("Chapter-15/guests.txt","r")
doc = docx.Document()

for name in nameTxt.readlines():

    para = doc.add_paragraph()
    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = para.add_run(
        'it would be a pleasure to have the company of \n').italics = True
    r = para.add_run(name)
    r = para.add_run("at 11010 Memory lane on the evening of \n").italics = True
    r = para.add_run("Apirl 1st\n")
    r = para.add_run("at 7 o'clock\n").italics = True
    doc.add_page_break()

doc.save("Chapter-15/invitations.docx")

